import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PRESET = "compact_reference_guide"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "667085"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
BODY_FONT = "Calibri"
CHINESE_FONT = "Microsoft YaHei"


def set_run_font(run, size=None, color=None, bold=None, italic=None, font=BODY_FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_style(style, size, color, before, after, line_spacing, bold=False):
    style.font.name = BODY_FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BODY_FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BODY_FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.widow_control = True


def add_numbering(document, num_id, abstract_id, fmt, text, left=540, hanging=270):
    numbering = document.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    level.append(lvl_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    ppr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), str(left))
    indent.set(qn("w:hanging"), str(hanging))
    ppr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    ppr.append(num_pr)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_geometry(table, widths):
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_tr_pr.append(repeat)


def widths_for_columns(count):
    patterns = {
        2: [2700, 6660],
        3: [2100, 2800, 4460],
        4: [1700, 2300, 2300, 3060],
        5: [1300, 1800, 1800, 2100, 2360]
    }
    if count in patterns:
        return patterns[count]
    base = CONTENT_WIDTH_DXA // count
    widths = [base] * count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(document, rows):
    column_count = len(rows[0])
    table = document.add_table(rows=1, cols=column_count)
    table.style = "Table Grid"
    for index, value in enumerate(rows[0]):
        cell = table.rows[0].cells[index]
        cell.text = value
        set_cell_shading(cell, HEADER_FILL)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                set_run_font(run, size=9.5, color=NAVY, bold=True)
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                if index == 0 and len(value) < 24:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, size=9.2, color="344054")
    set_table_geometry(table, widths_for_columns(column_count))
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_callout(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    ppr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CALLOUT_FILL)
    ppr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), BLUE)
    left.set(qn("w:space"), "8")
    borders.append(left)
    ppr.append(borders)
    run = paragraph.add_run(text)
    set_run_font(run, size=10, color="475467")


def add_code_block(document, lines):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.05
    ppr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F4F7")
    ppr.append(shading)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, size=8.4, color="344054", font="Consolas")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CHINESE_FONT)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9, color=MUTED)


def setup_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    configure_style(normal, 11, "222222", 0, 6, 1.25)
    configure_style(document.styles["Heading 1"], 16, BLUE, 18, 10, 1.0, True)
    configure_style(document.styles["Heading 2"], 13, BLUE, 14, 7, 1.0, True)
    configure_style(document.styles["Heading 3"], 12, DARK_BLUE, 10, 5, 1.0, True)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        document.styles[style_name].paragraph_format.keep_with_next = True
    title_style = document.styles.add_style("Knowledge Title", 1)
    configure_style(title_style, 24, NAVY, 10, 8, 1.0, True)
    title_style.paragraph_format.keep_with_next = True
    code_style = document.styles.add_style("Knowledge Code", 1)
    configure_style(code_style, 8.4, "344054", 0, 6, 1.05)
    add_numbering(document, 41, 41, "decimal", "%1.")
    add_numbering(document, 42, 42, "bullet", "•")
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    header_run = header.add_run("知识库 Agent 设计说明  ·  MVP 产品与技术参考")
    set_run_font(header_run, size=8.5, color=MUTED)
    add_page_field(section.footer.paragraphs[0])
    document.core_properties.title = "知识库 Agent 设计说明"
    document.core_properties.subject = "电商广告素材生成知识资产召回、过滤、排序与缺失处理"
    document.core_properties.author = "创意智造台产品团队"


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in raw):
            rows.append(raw)
        index += 1
    return rows, index


def build_docx(source_path, output_path):
    lines = source_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    setup_document(document)
    index = 0
    in_code = False
    code_lines = []
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and re.search(r"\|\s*:?-{3,}", lines[index + 1]):
            table_rows, index = parse_table(lines, index)
            add_markdown_table(document, table_rows)
            continue
        if stripped.startswith("# "):
            paragraph = document.add_paragraph(style="Knowledge Title")
            run = paragraph.add_run(stripped[2:])
            set_run_font(run, size=24, color=NAVY, bold=True)
        elif stripped.startswith("## "):
            document.add_paragraph(stripped[3:], style="Heading 1")
        elif stripped.startswith("### "):
            document.add_paragraph(stripped[4:], style="Heading 2")
        elif stripped.startswith("#### "):
            document.add_paragraph(stripped[5:], style="Heading 3")
        elif stripped.startswith("> "):
            add_callout(document, stripped[2:])
        elif re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            paragraph = document.add_paragraph(text)
            apply_numbering(paragraph, 41)
        elif stripped.startswith("- "):
            paragraph = document.add_paragraph(stripped[2:])
            apply_numbering(paragraph, 42)
        else:
            document.add_paragraph(stripped)
        index += 1
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    if not document.tables:
        raise RuntimeError("Document audit failed: no tables were generated")
    print(f"Built {output_path} with preset={PRESET}, tables={len(document.tables)}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build_docx(args.source.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
